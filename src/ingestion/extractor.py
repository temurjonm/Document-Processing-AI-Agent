import fitz
from docx import Document
from PIL import Image
import pytesseract
import os

def extract_text(file_path) -> list[dict]:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    elif ext in [".png", ".jpg", ".jpeg"]:
        return extract_text_from_image(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")   
    
def extract_text_from_pdf(file_path) -> list[dict]:
    pages = []
    doc = fitz.open(file_path)
    for i in range(len(doc)):
        page = doc.load_page(i)
        text = page.get_text()
        if len(text.strip()) < 50:
            text = _ocr_page(page)

        if text.strip():
            pages.append({"page": i + 1, "text": text})
    doc.close()
    return pages

def _ocr_page(page) -> str:
    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    img = img.convert("L")
    text = pytesseract.image_to_string(img)
    return text

def extract_text_from_docx(file_path) -> list[dict]:
    doc = Document(file_path)
    full_text = []
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            full_text.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                full_text.append(row_text)
    return [{"page": 1, "text": "\n".join(full_text)}]

def extract_text_from_image(file_path) -> list[dict]:
    img = Image.open(file_path)
    img = img.convert("L")
    text = pytesseract.image_to_string(img)
    return [{"page": 1, "text": text}]